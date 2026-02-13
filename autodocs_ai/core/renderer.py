"""Document rendering engines — Typst and LaTeX to PDF."""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from pathlib import Path

from autodocs_ai.config import RendererName


class RenderError(Exception):
    """Raised when document rendering fails."""


def render_typst(source: str, output_path: Path) -> Path:
    """Render Typst source to PDF.

    Args:
        source: Typst markup source code.
        output_path: Path for the output PDF file.

    Returns:
        Path to the generated PDF.

    Raises:
        RenderError: If Typst compilation fails.
    """
    # Try Python typst bindings first
    try:
        import typst as typst_lib

        with tempfile.NamedTemporaryFile(suffix=".typ", mode="w", delete=False) as f:
            f.write(source)
            f.flush()
            typ_path = Path(f.name)

        try:
            pdf_bytes = typst_lib.compile(typ_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_bytes(pdf_bytes)
            return output_path
        finally:
            typ_path.unlink(missing_ok=True)
    except ImportError:
        pass

    # Fall back to typst CLI
    typst_bin = shutil.which("typst")
    if not typst_bin:
        raise RenderError(
            "Typst is not installed. Install with: autodocs setup\n"
            "Or install the Python bindings: pip install typst"
        )

    with tempfile.NamedTemporaryFile(suffix=".typ", mode="w", delete=False) as f:
        f.write(source)
        f.flush()
        typ_path = Path(f.name)

    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        result = subprocess.run(
            [typst_bin, "compile", str(typ_path), str(output_path)],
            capture_output=True,
            text=True,
            timeout=60,
        )
        if result.returncode != 0:
            raise RenderError(f"Typst compilation failed:\n{result.stderr}")
        return output_path
    finally:
        typ_path.unlink(missing_ok=True)


def render_latex(source: str, output_path: Path) -> Path:
    """Render LaTeX source to PDF.

    Args:
        source: LaTeX source code.
        output_path: Path for the output PDF file.

    Returns:
        Path to the generated PDF.

    Raises:
        RenderError: If LaTeX compilation fails.
    """
    pdflatex = shutil.which("pdflatex") or shutil.which("xelatex")
    if not pdflatex:
        raise RenderError(
            "LaTeX is not installed. Install with: autodocs setup --latex\n"
            "Or install TinyTeX manually: https://yihui.org/tinytex/"
        )

    with tempfile.TemporaryDirectory() as tmpdir:
        tex_path = Path(tmpdir) / "document.tex"
        tex_path.write_text(source)

        # Run twice for cross-references
        for _ in range(2):
            result = subprocess.run(
                [pdflatex, "-interaction=nonstopmode", "-output-directory", tmpdir, str(tex_path)],
                capture_output=True,
                text=True,
                timeout=120,
            )

        pdf_path = Path(tmpdir) / "document.pdf"
        if not pdf_path.exists():
            raise RenderError(f"LaTeX compilation failed:\n{result.stderr}")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(pdf_path, output_path)
        return output_path


def render_html(source: str, output_path: Path) -> Path:
    """Write HTML source to file.

    Args:
        source: HTML source code.
        output_path: Path for the output HTML file.

    Returns:
        Path to the generated HTML file.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(source)
    return output_path


def render_markdown(source: str, output_path: Path) -> Path:
    """Write Markdown source to file.

    Args:
        source: Markdown source code.
        output_path: Path for the output Markdown file.

    Returns:
        Path to the generated Markdown file.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(source)
    return output_path


def render_docx(source: str, output_path: Path) -> Path:
    """Convert Markdown source to DOCX.

    Args:
        source: Markdown source (will be converted to DOCX).
        output_path: Path for the output DOCX file.

    Returns:
        Path to the generated DOCX file.

    Raises:
        RenderError: If DOCX generation fails.
    """
    try:
        from docx import Document
    except ImportError:
        raise RenderError(
            "python-docx is required for DOCX output. "
            "Install with: pip install autodocs-ai[extractors]"
        )

    doc = Document()
    lines = source.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped[0:3].rstrip(".").isdigit() and ". " in stripped:
            text = stripped.split(". ", 1)[1]
            doc.add_paragraph(text, style="List Number")
        elif stripped.startswith("---"):
            doc.add_page_break()
        else:
            doc.add_paragraph(stripped)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def render(
    source: str,
    output_path: Path,
    renderer: RendererName = RendererName.TYPST,
    output_format: str = "pdf",
) -> Path:
    """Render source content to the specified output format.

    Args:
        source: The source content to render.
        output_path: Path for the output file.
        renderer: Which rendering engine to use (typst or latex).
        output_format: Output format (pdf, docx, html, markdown).

    Returns:
        Path to the generated file.
    """
    if output_format == "html":
        return render_html(source, output_path)
    elif output_format == "markdown":
        return render_markdown(source, output_path)
    elif output_format == "docx":
        return render_docx(source, output_path)
    elif output_format == "pdf":
        if renderer == RendererName.LATEX:
            return render_latex(source, output_path)
        else:
            return render_typst(source, output_path)
    else:
        raise RenderError(f"Unsupported output format: {output_format}")
